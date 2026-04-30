"""
commercial_invoice.py — Builds the Commercial Invoice .docx document.

Entry point: create_commercial_invoice(invoice_data, assets_dir, output_dir)
Returns the full path to the saved .docx file.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.generators.base import (
    CONTENT_W, P_MARKS, P_DESC, P_QTY, P_PRICE, P_TOTAL,
    setup_page, build_header, build_info_block, port_delivery_details,
    build_payment_through, build_container_table, build_signature,
    add_small_spacer, replace_placeholders, prepare_placeholder_dict,
    set_col_width, thin_all_borders, set_cell_margins, clear_cell_paragraphs,
    set_paragraph_spacing, apply_run_format, set_table_width, remove_table_borders,
    set_cell_border, first_para, no_borders,
)
from storage.file_storage import save_docx
from app.config import COMMERCIAL_INVOICE_FILENAME


def build_invoice_title_row(doc):
    """'COMMERCIAL INVOICE' heading + Invoice No / Date row with thick bottom border."""
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(h, before=0, after=60)
    apply_run_format(h.add_run('COMMERCIAL INVOICE'), bold=True, size_pt=18)

    ref_table = doc.add_table(rows=1, cols=2)
    ref_table.autofit   = False
    ref_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(ref_table)
    set_table_width(ref_table, CONTENT_W)

    left  = ref_table.cell(0, 0)
    right = ref_table.cell(0, 1)
    set_col_width(left,  5.0)
    set_col_width(right, 2.5)
    set_cell_margins(left,  top=40, start=0,  bottom=40, end=80)
    set_cell_margins(right, top=40, start=80, bottom=40, end=0)

    thick_bottom = {'val': 'single', 'sz': 24, 'color': '000000'}
    none_border  = {'val': 'none',   'sz': 0,  'color': 'FFFFFF'}
    for cell in (left, right):
        set_cell_border(cell,
            top=none_border, bottom=thick_bottom,
            left=none_border, right=none_border)

    lp = first_para(left)
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(lp)
    apply_run_format(lp.add_run('REF: INVOICE NO:  '), bold=True,  size_pt=12)
    apply_run_format(lp.add_run('{invoice_no}'),       bold=False, size_pt=12)

    rp = first_para(right)
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(rp)
    apply_run_format(rp.add_run('DATE:  '),         bold=True,  size_pt=12)
    apply_run_format(rp.add_run('{invoice_date}'),  bold=False, size_pt=12)


def build_product_table(doc):
    """5-column product description table."""
    col_widths = [P_MARKS, P_DESC, P_QTY, P_PRICE, P_TOTAL]
    headers    = [
        'MARKS & NOS.', 'DESCRIPTION OF GOODS',
        'QTY./MT.', 'UNIT PRICE IN USD', 'TOTAL AMOUNT IN USD',
    ]

    t = doc.add_table(rows=3, cols=5)
    t.autofit   = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(t, CONTENT_W)

    def fmt(row_idx, col_idx, text, bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=10):
        c = t.cell(row_idx, col_idx)
        set_col_width(c, col_widths[col_idx].inches)
        thin_all_borders(c)
        set_cell_margins(c, top=40, start=60, bottom=40, end=60)
        clear_cell_paragraphs(c)
        p = c.add_paragraph()
        p.alignment = align
        set_paragraph_spacing(p)
        apply_run_format(p.add_run(text), bold=bold, size_pt=size_pt)

    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    RIGHT  = WD_ALIGN_PARAGRAPH.RIGHT

    # Row 0 — headers
    for col, hdr in enumerate(headers):
        fmt(0, col, hdr, bold=True, align=CENTER)

    # Row 1 — data
    fmt(1, 0, '{marks}')
    fmt(1, 1, '{description}')
    fmt(1, 2, '{qty_mt}',       align=CENTER)
    fmt(1, 3, '{unit_price}',   align=CENTER)
    fmt(1, 4, '{total_amount}', align=CENTER)

    # Row 2 — totals
    fmt(2, 0, '')
    fmt(2, 1, 'TOTAL :',         bold=True, align=RIGHT)
    fmt(2, 2, '{qty_mt}',        bold=True, align=CENTER)
    fmt(2, 3, 'TOTAL CFR VALUE', bold=True, align=CENTER)
    fmt(2, 4, '$ {total_amount}',bold=True, align=CENTER)


def build_amount_in_words(doc):
    """Full-width USD IN WORD line."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    remove_table_borders(t)
    set_table_width(t, CONTENT_W)

    c = t.cell(0, 0)
    set_col_width(c, 7.5)
    no_borders(c)
    set_cell_margins(c, top=40, start=0, bottom=40, end=0)
    clear_cell_paragraphs(c)

    p = c.add_paragraph()
    set_paragraph_spacing(p)
    apply_run_format(p.add_run('AMOUNT IN WORD: '),        bold=True, size_pt=11)
    apply_run_format(p.add_run('{amount_in_words}'),    size_pt=11)


# ── Entry Point ───────────────────────────────────────────────────────────────

def create_commercial_invoice(invoice_data: dict,
                               assets_dir: str,
                               output_dir: str) -> str:
    """
    Build and save the Commercial Invoice .docx.

    Args:
        invoice_data : extracted invoice dict (from InvoiceService)
        assets_dir   : path to folder containing PNG assets
        output_dir   : path to folder where .docx will be saved

    Returns:
        Full path to the saved .docx file.
    """
    doc = Document()
    setup_page(doc)
    build_header(doc, assets_dir)
    build_invoice_title_row(doc)

    add_small_spacer(doc, 4)
    build_info_block(doc)

    port_delivery_details(doc)
    add_small_spacer(doc, 4)

    build_payment_through(doc)
    add_small_spacer(doc, 4)

    build_product_table(doc)
    add_small_spacer(doc, 4)

    build_amount_in_words(doc)
    add_small_spacer(doc, 4)

    build_container_table(doc)
    add_small_spacer(doc, 12)

    build_signature(doc)

    replace_placeholders(doc, prepare_placeholder_dict(invoice_data))

    return save_docx(doc, output_dir, COMMERCIAL_INVOICE_FILENAME)
