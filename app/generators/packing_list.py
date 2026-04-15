"""
packing_list.py — Builds the Packing List .docx document.

Entry point: create_packing_list(invoice_data, assets_dir, output_dir)
Returns the full path to the saved .docx file.
"""

import math
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

from app.generators.base import (
    CONTENT_W,
    setup_page, build_header, build_info_block, port_delivery_details,
    build_payment_through, build_container_table, build_signature,
    add_small_spacer, replace_placeholders, prepare_placeholder_dict,
    set_col_width, thin_all_borders, set_cell_margins, clear_cell_paragraphs,
    set_paragraph_spacing, apply_run_format, set_table_width, remove_table_borders,
    set_cell_border, first_para, no_borders,
)
from storage.file_storage import save_docx
from app.config import PACKING_LIST_FILENAME

# ── Packing List Layout Constants ─────────────────────────────────────────────
HALF_W    = 3.6    # inches — each half-table width
GAP_W     = 0.3    # inches — gap between the two halves

PL_PALLET = 0.65
PL_SPOOL  = 1.0
PL_GROSS  = 0.975
PL_NET    = 0.975
# Sum: 0.65 + 1.0 + 0.975 + 0.975 = 3.6"

PL_COL_WIDTHS = [PL_PALLET, PL_SPOOL, PL_GROSS, PL_NET]
PL_HEADERS    = ['Pallet No.', 'No. of Spool', 'Gross Wt.\n(Kgs)', 'Net Wt.\n(Kgs)']


# ── Section Builders ──────────────────────────────────────────────────────────

def build_packing_list_title_row(doc):
    """'PACKING LIST' heading + Invoice No / Date row with thick bottom border."""
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(h, before=0, after=60)
    apply_run_format(h.add_run('PACKING LIST'), bold=True, size_pt=18)

    ref_table = doc.add_table(rows=1, cols=2)
    ref_table.autofit   = False
    ref_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(ref_table)
    set_table_width(ref_table, CONTENT_W)

    left  = ref_table.cell(0, 0)
    right = ref_table.cell(0, 1)
    set_col_width(left,  5.0)
    set_col_width(right, 2.5)
    set_cell_margins(left,  top=40, start=0,  bottom=40, end=80)
    set_cell_margins(right, top=40, start=80, bottom=40, end=0)

    thick_bottom = {'val': 'single', 'sz': 24, 'color': '000000'}
    none_border  = {'val': 'none',   'sz': 0,  'color': 'FFFFFF'}
    for cell in (left, right):
        set_cell_border(cell,
            top=none_border, bottom=thick_bottom,
            left=none_border, right=none_border)

    lp = first_para(left)
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(lp)
    apply_run_format(lp.add_run('REF: INVOICE NO:  '), bold=True,  size_pt=12)
    apply_run_format(lp.add_run('{invoice_no}'),       bold=False, size_pt=12)

    rp = first_para(right)
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(rp)
    apply_run_format(rp.add_run('DATE:  '),        bold=True,  size_pt=12)
    apply_run_format(rp.add_run('{invoice_date}'), bold=False, size_pt=12)


def build_pl_product_table(doc):
    """2-column product table for Packing List: MARKS & NOS | DESCRIPTION OF GOODS."""
    col_widths = [2.0, 5.5]
    headers    = ['MARKS & NOS.', 'DESCRIPTION OF GOODS']

    t = doc.add_table(rows=2, cols=2)
    t.autofit   = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(t, CONTENT_W)

    def fmt(row_idx, col_idx, text, bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=10):
        c = t.cell(row_idx, col_idx)
        set_col_width(c, col_widths[col_idx])
        thin_all_borders(c)
        set_cell_margins(c, top=40, start=60, bottom=40, end=60)
        clear_cell_paragraphs(c)
        p = c.add_paragraph()
        p.alignment = align
        set_paragraph_spacing(p)
        apply_run_format(p.add_run(text), bold=bold, size_pt=size_pt)

    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    for col, hdr in enumerate(headers):
        fmt(0, col, hdr, bold=True, align=CENTER)
    fmt(1, 0, '{marks}')
    fmt(1, 1, '{description}')


def _build_packing_list_half_table(outer_cell, rows_data: list):
    """Build one bordered 4-column packing list table inside an outer wrapper cell."""
    total_rows = 1 + len(rows_data) + 1

    t = outer_cell.add_table(rows=total_rows, cols=4)
    t.autofit = False
    set_table_width(t, Inches(HALF_W))

    def fmt(row_idx, col_idx, text, bold=False,
            align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=9):
        c = t.cell(row_idx, col_idx)
        set_col_width(c, PL_COL_WIDTHS[col_idx])
        thin_all_borders(c)
        set_cell_margins(c, top=30, start=40, bottom=30, end=40)
        clear_cell_paragraphs(c)
        p = c.add_paragraph()
        p.alignment = align
        set_paragraph_spacing(p)
        apply_run_format(p.add_run(str(text)), bold=bold, size_pt=size_pt)

    # Header row
    for col, hdr in enumerate(PL_HEADERS):
        fmt(0, col, hdr, bold=True)

    # Data rows
    total_spool = 0
    total_gross = 0.0
    total_net   = 0.0

    for i, row in enumerate(rows_data):
        ri = i + 1
        fmt(ri, 0, row['pallet_no'])
        fmt(ri, 1, row['no_of_spool'])
        fmt(ri, 2, f"{row['gross_weight_kgs']:,.2f}")
        fmt(ri, 3, f"{row['net_weight_kgs']:,.2f}")
        total_spool += row['no_of_spool']
        total_gross += row['gross_weight_kgs']
        total_net   += row['net_weight_kgs']

    # Totals row
    tr = 1 + len(rows_data)
    fmt(tr, 0, 'TOTAL',                    bold=True)
    fmt(tr, 1, str(total_spool),            bold=True)
    fmt(tr, 2, f"{total_gross:,.2f}",       bold=True)
    fmt(tr, 3, f"{total_net:,.2f}",         bold=True)


def build_packing_list_table(doc, packing_list: list):
    """Side-by-side packing list: left half (first ceil(N/2)) | gap | right half (rest)."""
    n     = len(packing_list)
    split = math.ceil(n / 2)
    left_rows  = packing_list[:split]
    right_rows = packing_list[split:]

    outer = doc.add_table(rows=1, cols=3)
    outer.autofit   = False
    outer.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(outer)
    set_table_width(outer, CONTENT_W)

    left_cell  = outer.cell(0, 0)
    gap_cell   = outer.cell(0, 1)
    right_cell = outer.cell(0, 2)

    set_col_width(left_cell,  HALF_W)
    set_col_width(gap_cell,   GAP_W)
    set_col_width(right_cell, HALF_W)

    for cell in (left_cell, gap_cell, right_cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        no_borders(cell)
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        clear_cell_paragraphs(cell)

    gap_cell.add_paragraph()

    _build_packing_list_half_table(left_cell,  left_rows)
    _build_packing_list_half_table(right_cell, right_rows)


# ── Entry Point ───────────────────────────────────────────────────────────────

def create_packing_list(invoice_data: dict,
                         assets_dir: str,
                         output_dir: str) -> str:
    """
    Build and save the Packing List .docx.

    Args:
        invoice_data : extracted invoice dict (from InvoiceService)
        assets_dir   : path to folder containing PNG assets
        output_dir   : path to folder where .docx will be saved

    Returns:
        Full path to the saved .docx file.
    """
    doc = Document()
    setup_page(doc)
    build_header(doc, assets_dir)
    build_packing_list_title_row(doc)

    add_small_spacer(doc, 4)
    build_info_block(doc)

    port_delivery_details(doc)
    add_small_spacer(doc, 4)

    build_payment_through(doc)
    add_small_spacer(doc, 4)

    build_pl_product_table(doc)
    add_small_spacer(doc, 4)

    packing_list = invoice_data.get('packing_list_details', {}).get('packing_list', [])
    build_packing_list_table(doc, packing_list)
    add_small_spacer(doc, 4)

    build_container_table(doc)
    add_small_spacer(doc, 12)

    build_signature(doc)

    replace_placeholders(doc, prepare_placeholder_dict(invoice_data))

    return save_docx(doc, output_dir, PACKING_LIST_FILENAME)
