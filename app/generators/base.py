"""
base.py — Shared document building blocks used by both generators.

All XML helpers, font constants, page setup, and paragraph utilities live here.
commercial_invoice.py and packing_list.py import from this module.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Font & Layout Constants ───────────────────────────────────────────────────
FONT_NAME   = 'Times New Roman'
CONTENT_W   = Inches(7.5)

# Info block column widths
LEFT_COL_W  = Inches(4.0)
RIGHT_COL_W = Inches(3.5)

# Product table column widths (sum = 7.5")
P_MARKS = Inches(1.4)
P_DESC  = Inches(3.1)
P_QTY   = Inches(0.8)
P_PRICE = Inches(1.1)
P_TOTAL = Inches(1.1)

# Container table column widths (sum = 7.5")
C_CNUM  = Inches(1.5)
C_SEAL  = Inches(1.3)
C_SIZE  = Inches(0.9)
C_PLTS  = Inches(1.3)
C_GROSS = Inches(1.25)
C_NET   = Inches(1.25)


# ── XML Cell / Border Helpers ─────────────────────────────────────────────────

def set_cell_margins(cell, top=40, start=0, bottom=40, end=0):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(existing)
    tcMar = OxmlElement('w:tcMar')
    for name, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side, attrs in kwargs.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   attrs.get('val',   'single'))
        el.set(qn('w:sz'),    str(attrs.get('sz', 4)))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), attrs.get('color', '000000'))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def no_borders(cell):
    set_cell_border(cell,
        top    = {'val': 'none', 'sz': 0, 'color': 'FFFFFF'},
        bottom = {'val': 'none', 'sz': 0, 'color': 'FFFFFF'},
        left   = {'val': 'none', 'sz': 0, 'color': 'FFFFFF'},
        right  = {'val': 'none', 'sz': 0, 'color': 'FFFFFF'},
    )


def thin_all_borders(cell):
    b = {'val': 'single', 'sz': 4, 'color': '000000'}
    set_cell_border(cell, top=b, bottom=b, left=b, right=b)


def remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_table_width(table, width_emu):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    dxa  = int(width_emu * 1440 / 914400)
    tblW.set(qn('w:w'),    str(dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def set_col_width(cell, width_inches):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# ── Paragraph Helpers ─────────────────────────────────────────────────────────

def set_paragraph_spacing(para, before=0, after=0, line=240):
    pPr = para._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:spacing')):
        pPr.remove(existing)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'),   str(before))
    spacing.set(qn('w:after'),    str(after))
    spacing.set(qn('w:line'),     str(line))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def apply_run_format(run, bold=False, size_pt=9, italic=False):
    run.font.name      = FONT_NAME
    run.font.size      = Pt(size_pt)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def clear_cell_paragraphs(cell):
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)


def first_para(cell):
    p = cell.paragraphs[0]
    p.clear()
    set_paragraph_spacing(p)
    return p


def add_small_spacer(doc, height_pt=4):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line=int(height_pt * 20))


# ── Page Setup ────────────────────────────────────────────────────────────────

def setup_page(doc):
    section = doc.sections[0]
    section.page_width      = Inches(8.5)
    section.page_height     = Inches(11)
    section.left_margin     = Inches(0.5)
    section.right_margin    = Inches(0.5)
    section.top_margin      = Inches(1.0)
    section.bottom_margin   = Inches(0.4)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.2)


# ── Shared Section Builders ───────────────────────────────────────────────────

def build_header(doc, assets_dir: str):
    """Letterhead header: logo | contact info, address line, slogan."""
    logo_path   = os.path.join(assets_dir, "JJF letter_head.png")
    slogan_path = os.path.join(assets_dir, "JJF header bottom line.png")

    section = doc.sections[0]
    header  = section.header
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    logo_table = header.add_table(rows=1, cols=2, width=CONTENT_W)
    logo_table.autofit    = False
    logo_table.alignment  = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(logo_table)

    left_cell  = logo_table.cell(0, 0)
    right_cell = logo_table.cell(0, 1)
    set_col_width(left_cell,  4.9)
    set_col_width(right_cell, 2.6)
    left_cell.vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    no_borders(left_cell)
    no_borders(right_cell)
    set_cell_margins(left_cell,  top=0,  start=0,  bottom=0,  end=80)
    set_cell_margins(right_cell, top=40, start=80, bottom=40, end=0)

    lp = first_para(left_cell)
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(lp)
    if os.path.exists(logo_path):
        lp.add_run().add_picture(logo_path, width=Inches(4.9))
    else:
        r = lp.add_run('JUNAID JUTE FIBERS')
        apply_run_format(r, bold=True, size_pt=18)

    contact_lines = [
        'Phone: +8801743108161',
        'E-mail: sazzad_zahangir@mail.ru',
        'Dhaka Office: 70/5, Zigatola Post Office Road.',
        'Zigatola, Dhaka-1209, Bangladesh.',
    ]
    rp = first_para(right_cell)
    rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(rp)
    for i, line in enumerate(contact_lines):
        r = rp.add_run(line)
        apply_run_format(r, size_pt=8.5)
        if i < len(contact_lines) - 1:
            rp.add_run().add_break()

    addr_para = header.add_paragraph()
    set_paragraph_spacing(addr_para, before=20, after=0)
    r_addr = addr_para.add_run(
        'Address: HO: 160, Sher A Bangla Road, Khulna, Bangladesh.  License No: 1424'
    )
    apply_run_format(r_addr, bold=True, size_pt=10)

    slogan_para = header.add_paragraph()
    slogan_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(slogan_para)
    if os.path.exists(slogan_path):
        slogan_para.add_run().add_picture(slogan_path, width=CONTENT_W)
    else:
        r_s = slogan_para.add_run(
            '100% Export High Quality Jute Yarn, Sacking Bag, Hessian Cloth and Jute Rope'
        )
        apply_run_format(r_s, size_pt=8, italic=True)


def build_info_block(doc):
    """Outer 2-col wrapper: Shipper info (left) | Notify Party (right)."""
    outer = doc.add_table(rows=1, cols=2)
    outer.autofit   = False
    outer.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(outer)
    set_table_width(outer, CONTENT_W)

    left_cell  = outer.cell(0, 0)
    right_cell = outer.cell(0, 1)
    set_col_width(left_cell,  LEFT_COL_W.inches)
    set_col_width(right_cell, RIGHT_COL_W.inches)
    left_cell.vertical_alignment  = WD_ALIGN_VERTICAL.TOP
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    no_borders(left_cell)
    no_borders(right_cell)
    set_cell_margins(left_cell,  top=0, start=0, bottom=0, end=0)
    set_cell_margins(right_cell, top=0, start=0, bottom=0, end=0)
    clear_cell_paragraphs(left_cell)
    clear_cell_paragraphs(right_cell)

    _build_left_inner_table(left_cell)
    _build_right_inner_table(right_cell)


def _build_left_inner_table(outer_left_cell):
    rows = [
        ('label',  'SELLER:',  ''),
        ('value',  '',                   '{seller_name}'),
        ('address','',                   '{seller_address_1} {seller_address_2}'),
        ('address', '',                   'LICENSE NO: {seller_license}'),
        ('spacer', '',                   ''),
        ('lv',     'EXP NO:  ',          '{exp_no}  DATE: {exp_date}'),
        ('lv',     'LC/SC/NO: ',         '{lc_sc_no}  DATE: {lc_sc_date}'),
        ('lv',     'BILL OF LADING: ',   '{bl_no}  DATE: {bl_date}'),
    ]
    t = outer_left_cell.add_table(rows=len(rows), cols=1)
    t.autofit = False
    remove_table_borders(t)
    set_table_width(t, LEFT_COL_W)

    for i, (kind, label, value) in enumerate(rows):
        c = t.cell(i, 0)
        set_col_width(c, LEFT_COL_W.inches)
        no_borders(c)
        set_cell_margins(c, top=80, start=0, bottom=0, end=40)
        clear_cell_paragraphs(c)

        if kind == 'label':
            p = c.add_paragraph()
            set_paragraph_spacing(p, before=40, after=0)
            apply_run_format(p.add_run(label), bold=True, size_pt=13)
        elif kind == 'value':
            p = c.add_paragraph()
            set_paragraph_spacing(p)
            apply_run_format(p.add_run(value), bold=True, size_pt=11)
        elif kind == 'address':
            p = c.add_paragraph()
            set_paragraph_spacing(p, before=10)
            apply_run_format(p.add_run(value), size_pt=9)
        elif kind == 'spacer':
            p = c.add_paragraph()
            set_paragraph_spacing(p, before=0, after=0, line=160)
        elif kind == 'lv':
            p = c.add_paragraph()
            set_paragraph_spacing(p)
            apply_run_format(p.add_run(label), bold=True, size_pt=11)
            apply_run_format(p.add_run(value), size_pt=11)


def _build_right_inner_table(outer_right_cell):
    rows = [
        ('label',  'NOTIFY PARTY:',       ''),
        ('value',  '',                     '{notify_name}'),
        ('address','',                     '{notify_address_1}{notify_address_2}'),
        ('spacer', '',                     ''),
        ('lv',     'H.S. CODE NO: ',      '{hs_code}'),
        ('lv',     'COUNTRY OF ORIGIN: ', '{country}'),
        ('lv',     'PI NO: ',             '{pi_no}'),
    ]
    t = outer_right_cell.add_table(rows=len(rows), cols=1)
    t.autofit = False
    remove_table_borders(t)
    set_table_width(t, RIGHT_COL_W)

    for i, (kind, label, value) in enumerate(rows):
        c = t.cell(i, 0)
        set_col_width(c, RIGHT_COL_W.inches)
        no_borders(c)
        set_cell_margins(c, top=80, start=40, bottom=0, end=0)
        clear_cell_paragraphs(c)

        if kind == 'label':
            p = c.add_paragraph()
            set_paragraph_spacing(p, before=40, after=0)
            apply_run_format(p.add_run(label), bold=True, size_pt=13)
        elif kind == 'value':
            p = c.add_paragraph()
            set_paragraph_spacing(p)
            apply_run_format(p.add_run(value), bold=True, size_pt=11)
        elif kind == 'address':
            p = c.add_paragraph()
            set_paragraph_spacing(p, before=10)
            apply_run_format(p.add_run(value), size_pt=9)
        elif kind == 'spacer':
            p = c.add_paragraph()
            set_paragraph_spacing(p, before=0, after=0, line=160)
        elif kind == 'lv':
            p = c.add_paragraph()
            set_paragraph_spacing(p)
            apply_run_format(p.add_run(label), bold=True, size_pt=11)
            apply_run_format(p.add_run(value), size_pt=11)


def port_delivery_details(doc):
    t = doc.add_table(rows=3, cols=1)
    t.autofit   = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(t)
    set_table_width(t, CONTENT_W)

    for row_idx, (label, placeholder) in enumerate([
        ('PORT OF DISCHARGE: ', '{port_discharge}'),
        ('PORT OF LOADING: ',   '{port_loading}'),
        ('TERMS OF DELIVERY: ', '{terms}'),
    ]):
        c = t.cell(row_idx, 0)
        set_col_width(c, 7.5)
        no_borders(c)
        set_cell_margins(c, top=80, start=0, bottom=0, end=0)
        clear_cell_paragraphs(c)
        p = c.add_paragraph()
        set_paragraph_spacing(p)
        apply_run_format(p.add_run(label),       bold=True, size_pt=11)
        apply_run_format(p.add_run(placeholder), size_pt=11)


def build_payment_through(doc):
    """PAYMENT THROUGH section with fixed seller bank information."""
    add_small_spacer(doc, 0)
    t = doc.add_table(rows=1, cols=2)
    t.autofit   = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(t)
    set_table_width(t, CONTENT_W)

    bank_text = (
        # 'Beneficiary Name: JUNAID JUTE FIBERS, '
        'ISLAMI BANK PLC, '
        'Savar Branch, Dhaka, Bangladesh, '
        'A/C No: 20501300100650709, '
        'SWIFT: IBBLBDDH130, '
        'Routing No: 125264097'
    )

    for col_idx, (width, text, bold) in enumerate([
        (1.7, 'PAYMENT THROUGH:', True),
        (5.8, bank_text, False),
    ]):
        c = t.cell(0, col_idx)
        set_col_width(c, width)
        no_borders(c)
        set_cell_margins(c, top=0, start=0, bottom=80, end=0)
        clear_cell_paragraphs(c)
        p = c.add_paragraph()
        set_paragraph_spacing(p)
        apply_run_format(p.add_run(text), bold=bold, size_pt=11)


def build_container_table(doc):
    col_widths = [C_CNUM, C_SEAL, C_SIZE, C_PLTS, C_GROSS, C_NET]
    headers    = [
        'CONTAINER NO.', 'SEAL NO', 'CONTAINER SIZE',
        'PALLETS/TRUSS',  'GROSS WT (KG)', 'NET WT (KG)',
    ]
    t = doc.add_table(rows=3, cols=6)
    t.autofit   = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(t, CONTENT_W)

    def fmt(row_idx, col_idx, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=9.5):
        c = t.cell(row_idx, col_idx)
        set_col_width(c, col_widths[col_idx].inches)
        thin_all_borders(c)
        set_cell_margins(c, top=40, start=60, bottom=40, end=60)
        clear_cell_paragraphs(c)
        p = c.add_paragraph()
        p.alignment = align
        set_paragraph_spacing(p)
        apply_run_format(p.add_run(text), bold=bold, size_pt=size_pt)

    for col, hdr in enumerate(headers):
        fmt(0, col, hdr, bold=True)

    for col, val in enumerate([
        '{container_no}', '{seal_no}', '{container_size}',
        '{pallets}',       '{gross_wt}', '{net_wt}'
    ]):
        fmt(1, col, val)

    fmt(2, 0, '')
    fmt(2, 1, '')
    fmt(2, 2, '')
    fmt(2, 3, 'TOTAL',         bold=True)
    fmt(2, 4, '{total_gross}', bold=True)
    fmt(2, 5, '{total_net}',   bold=True)


def build_signature(doc):
    t = doc.add_table(rows=1, cols=2)
    t.autofit   = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(t)
    set_table_width(t, CONTENT_W)

    left  = t.cell(0, 0)
    right = t.cell(0, 1)
    set_col_width(left,  4.5)
    set_col_width(right, 3.0)
    no_borders(left)
    no_borders(right)
    set_cell_margins(left,  top=20, start=0,  bottom=0, end=0)
    set_cell_margins(right, top=20, start=80, bottom=0, end=0)
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    clear_cell_paragraphs(left)
    clear_cell_paragraphs(right)

    set_paragraph_spacing(left.add_paragraph())

    for text, bold in [('FOR & ON BEHALF OF', True), ('JUNAID JUTE FIBER', True)]:
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p)
        apply_run_format(p.add_run(text), bold=bold, size_pt=9)

    for _ in range(3):
        p = right.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line=200)

    p_sig = right.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_sig, before=0, after=0)

    pPr  = p_sig._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '000000')
    pBdr.append(top)
    pPr.append(pBdr)

    apply_run_format(p_sig.add_run('(Md. Sazzad Zahangir)'), size_pt=8.5)


# ── Placeholder Replacement ───────────────────────────────────────────────────

def replace_placeholders(doc, data_dict: dict):
    """Replace {placeholder} tokens throughout the document (body + tables + header)."""

    def replace_in_para(para):
        for key, value in data_dict.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))

    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)
                for nested in cell.tables:
                    replace_in_table(nested)

    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        replace_in_table(table)
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_para(para)
        for table in section.header.tables:
            replace_in_table(table)


def prepare_placeholder_dict(json_data: dict) -> dict:
    """Build the {placeholder}: value mapping from invoice JSON data."""
    container = json_data.get("containers", [{}])[0]
    summary   = json_data.get("packing_list_details", {}).get("summary", {})

    return {
        "{invoice_no}":       json_data.get("invoice_details", {}).get("invoice_no", ""),
        "{invoice_date}":     json_data.get("invoice_details", {}).get("invoice_date", ""),
        "{exp_no}":           json_data.get("invoice_details", {}).get("exp_no", ""),
        "{exp_date}":         json_data.get("invoice_details", {}).get("exp_date", ""),
        "{lc_sc_no}":         json_data.get("invoice_details", {}).get("lc_sc_no", ""),
        "{lc_sc_date}":       json_data.get("invoice_details", {}).get("lc_sc_date", ""),
        "{bl_no}":            json_data.get("invoice_details", {}).get("bl_no", ""),
        "{bl_date}":          json_data.get("invoice_details", {}).get("bl_date", ""),
        "{pi_no}":            json_data.get("invoice_details", {}).get("pi_no", ""),
        "{shipper_name}":     json_data.get("shipper", {}).get("name", ""),
        "{shipper_address_1}":json_data.get("shipper", {}).get("address_line1", ""),
        "{shipper_address_2}":json_data.get("shipper", {}).get("address_line2", ""),
        "{notify_name}":      json_data.get("notify_party", {}).get("name", ""),
        "{notify_address_1}": json_data.get("notify_party", {}).get("address_line1", ""),
        "{notify_address_2}": json_data.get("notify_party", {}).get("address_line2", ""),
        "{seller_name}":      json_data.get("for_account_of", {}).get("name", ""),
        "{seller_address_1}": json_data.get("for_account_of", {}).get("address_line1", ""),
        "{seller_address_2}": json_data.get("for_account_of", {}).get("address_line2", ""),
        "{seller_license}":   json_data.get("for_account_of", {}).get("license_no", ""),
        "{terms}":            json_data.get("invoice_details", {}).get("terms_of_delivery", ""),
        "{hs_code}":          json_data.get("invoice_details", {}).get("hs_code", ""),
        "{country}":          json_data.get("invoice_details", {}).get("country_of_origin", ""),
        "{port_loading}":     json_data.get("port_info", {}).get("port_of_loading", ""),
        "{port_discharge}":   json_data.get("port_info", {}).get("port_of_discharge", ""),
        "{marks}":            json_data.get("goods", {}).get("marks_and_nos", ""),
        "{description}":      json_data.get("goods", {}).get("description", ""),
        "{qty_mt}":           json_data.get("goods", {}).get("quantity_mt", 0),
        "{unit_price}":       json_data.get("goods", {}).get("unit_price_usd", 0),
        "{total_amount}":     json_data.get("goods", {}).get("total_amount_usd", 0),
        "{amount_in_words}":  json_data.get("amount_in_words", ""),
        "{container_no}":     container.get("container_no", ""),
        "{seal_no}":          container.get("seal_no", ""),
        "{container_size}":   container.get("container_size", ""),
        "{pallets}":          container.get("pallets_truss", 0),
        "{gross_wt}":         container.get("gross_weight_kg", 0),
        "{net_wt}":           container.get("net_weight_kg", 0),
        "{total_gross}":      summary.get("total_gross_weight_kgs", 0),
        "{total_net}":        summary.get("total_net_weight_kgs", 0),

    }
